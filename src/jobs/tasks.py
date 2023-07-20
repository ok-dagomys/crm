import asyncio
import logging
import os
from datetime import datetime
from time import sleep

import aspose.words as aw
import docx
import pandas as pd
from tqdm import tqdm

from cfg import task_source, task_registry, task_archive
from src.service.tasks import task_to_db

logging.basicConfig(level=logging.INFO)
file_list = []


def convert_to_docx(file_doc, c_time, m_time):
    file_docx = file_doc.rsplit('.', 1)[0] + '.docx'
    new_filename = os.path.join(task_source, file_docx)
    old_filename = os.path.join(task_source, file_doc)
    doc = aw.Document(old_filename)
    doc.save(new_filename)
    os.remove(f'{task_source}/{file_doc}')
    os.utime(f'{task_source}/{file_docx}', (c_time, m_time))
    return file_docx


def make_file_list(f_docx, c_date):
    doc = docx.Document(f'{task_source}/{f_docx}')
    cost = 'Стоимость не указана'
    for paragraph in doc.paragraphs:
        if 'Предположительная стоимость заявки:' in paragraph.text:
            cost = paragraph.text.split(':')[-1].strip()

    file_docx = os.path.splitext(f_docx)[0]
    if '-' in file_docx:
        status = 'заявка запланирована в работу'
    elif file_docx.count('+') == 1:
        status = 'заявка подписана и передана в мто'
    elif file_docx.count('+') == 2:
        status = 'подготовлен договор и счет'
    elif file_docx.count('+') == 3:
        status = 'счет подписан и передан в оплату'
    elif file_docx.count('+') == 4:
        status = 'товар поставлен'
    elif '=' in file_docx:
        status = 'заявка готовится к торгам'
    else:
        status = 'статус не присвоен'

    file_list.append([file_docx.lower(), c_date.strftime('%Y.%m.%d'), cost, status])


def write_xlsx(dataframe):
    with pd.ExcelWriter(f'{task_registry}/{datetime.now().year}.xlsx', engine='xlsxwriter') as wb:
        dataframe.to_excel(wb, sheet_name='Реестр', index=False)
        sheet = wb.sheets['Реестр']

        sheet.set_column('A:A', 50)
        sheet.set_column('B:B', 10)
        sheet.set_column('C:C', 20)
        sheet.set_column('D:D', 33)


def add_to_registry(f_list):
    df = pd.DataFrame(f_list, columns=['Заявка', 'Дата', 'Стоимость', 'Статус'])\
        .drop_duplicates()\
        .sort_values(by='Статус', ascending=True)\
        .dropna()\
        .reset_index(drop=True)

    if os.path.isfile(f'{task_registry}/{datetime.now().year}.xlsx'):
        n_list = pd.read_excel(f'{task_registry}/{datetime.now().year}.xlsx')\
                .dropna()\
                .values.tolist()

        if sorted(f_list) != sorted(n_list):
            write_xlsx(df)
    else:
        write_xlsx(df)


def transfer_to_archive(f_docx, f_date):
    if not os.path.isdir(f'{task_archive}/{f_date.year}'):
        os.mkdir(f'{task_archive}/{f_date.year}')
    os.replace(f'{task_source}/{f_docx}',
               f'{task_archive}/{f_date.year}/{f_date.strftime("%Y.%m.%d")} - {f_docx.split("+", 5)[-1].strip()}')
    logging.info(f' Task {f_docx.split("+", 5)[-1].strip()} completed')
    with tqdm(total=100) as pbar:
        for i in range(10):
            sleep(0.1)
            pbar.update(10)
            pbar.set_description("Transferring...")
    logging.info(f' Saved in archive as {f_date.strftime("%Y.%m.%d")} - {f_docx.split("+", 5)[-1].strip()}')
    task_to_db('new', f_docx.split("+", 5)[-1].strip(), 'Заявка исполнена и перенесена в архив')


async def scan_tasks():
    for file in os.listdir(task_source):
        created_time = os.path.getctime(f'{task_source}/{file}')
        modified_time = os.path.getmtime(f'{task_source}/{file}')
        date = datetime.fromtimestamp(modified_time)

        if file.lower().endswith('.doc') and '~' not in file:
            file = convert_to_docx(file, created_time, modified_time)
            make_file_list(file, date)

        elif file.lower().endswith('.docx') and '~' not in file:
            if file.count('+') < 5:
                make_file_list(file, date)
            elif file.count('+') >= 5:
                transfer_to_archive(file, date)

    add_to_registry(file_list)


async def check_task():
    task = asyncio.create_task(scan_tasks())
    await task
    logging.info(f' {datetime.now().strftime("%Y.%m.%d-%H:%M:%S")} | Tasks checked\n')
    await asyncio.sleep(0.1)


asyncio.run(check_task())
